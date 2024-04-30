import os
import openai
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
plt.switch_backend('agg')
from pandasai.llm.openai import OpenAI
from pandasai import SmartDataframe
from docx import Document
from docx.shared import Pt
from streamlit_chat import message
from dotenv import load_dotenv
import zipfile

load_dotenv()

# Loading API Key
openai.api_key = os.getenv('OPENAI_API_KEY')
llm = OpenAI(temperature=0)

# Hide traceback
st.set_option('client.showErrorDetails', False)

# Set page title and header
st.set_page_config(page_title="CHAT BOT", page_icon=":robot_face:")
st.markdown("<h1 style='text-align: center;'>DATA ANALYST</h1>", unsafe_allow_html=True)

# Initialise session state variables
if 'generated' not in st.session_state:
    st.session_state['generated'] = []
if 'past' not in st.session_state:
    st.session_state['past'] = []
if 'messages' not in st.session_state:
    st.session_state['messages'] = [{"role": "system", "content": "You are a helpful assistant."}]

# Allow user to upload file(s)
uploaded_files = st.file_uploader("**Choose file(s):**", accept_multiple_files=True)

# Process the uploaded file(s)
dataframes = []
uploaded_file_names = []
if uploaded_files:
    for uploaded_file in uploaded_files:
        file_extension = os.path.splitext(uploaded_file.name)[1]
        if file_extension.lower() == ".zip":
            # Create a temporary directory to extract the contents of the .zip file
            temp_dir = "temp_zip_extract"
            os.makedirs(temp_dir, exist_ok=True)
            
            # Extract the contents of the .zip file to the temporary directory
            with zipfile.ZipFile(uploaded_file, "r") as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Process each extracted file
            for filename in os.listdir(temp_dir):
                filepath = os.path.join(temp_dir, filename)
                if filename.endswith(".csv"):
                    dataframe = pd.read_csv(filepath)
                elif filename.endswith((".xls", ".xlsx", "xlsm", ".xlsb")):
                    dataframe = pd.read_excel(filepath)
                elif filename.endswith(".txt"):
                    # Check the second line of the file to determine the delimiter
                    with open(filepath, 'r') as f:
                        # Skip the first line
                        next(f)
                        second_line = f.readline().strip()
                        delimiter = "|" if "|" in second_line else "\t"  # Assume "|" delimiter for .psv, "\t" for .tsv

                        # Reset file pointer to the beginning of the file
                        f.seek(0)

                        # Read the file using pandas read_csv
                        dataframe = pd.read_csv(f, delimiter=delimiter)
                else:
                    st.warning(f"Unsupported file format: {filename}")
                    continue
                dataframes.append(dataframe)
                uploaded_file_names.append(filename)
            
            # Remove the temporary directory and its contents
            if os.path.exists(temp_dir):
                for root, dirs, files in os.walk(temp_dir, topdown=False):
                    for name in files:
                        os.remove(os.path.join(root, name))
                    for name in dirs:
                        os.rmdir(os.path.join(root, name))
                os.rmdir(temp_dir)
            
        elif file_extension.lower() in [".csv", ".xls", ".xlsx", "xlsm", ".xlsb", ".txt"]:
            # Process single file uploads as usual
            if file_extension.lower() == ".csv":
                dataframe = pd.read_csv(uploaded_file)
            elif file_extension.lower() in [".xls", ".xlsx", "xlsm", ".xlsb"]:
                dataframe = pd.read_excel(uploaded_file)
            elif file_extension.lower() == ".txt":
                # Check the second line of the file to determine the delimiter
                with uploaded_file as f:
                    # Skip the first line
                    next(f)
                    second_line = f.readline().decode().strip()  # decode bytes to string
                    delimiter = "|" if "|" in second_line else "\t"  # Assume "|" delimiter for .psv, "\t" for .tsv

                    # Reset file pointer to the beginning of the file
                    f.seek(0)

                    # Read the file using pandas read_csv
                    dataframe = pd.read_csv(f, delimiter=delimiter)
            dataframes.append(dataframe)
            uploaded_file_names.append(uploaded_file.name)
        else:
            st.warning(f"Unsupported file format: {file_extension}")

# Display the dataframes
if dataframes:
    for idx, dataframe in enumerate(dataframes):
        st.write(f"{idx+1}.{uploaded_file_names[idx]}")  # Display file name
        st.write(dataframe)

# Allow user to select file for chatting
if dataframes:
    selected_file_name = st.selectbox("**Select file for chatting:**", options=uploaded_file_names)

    # Find the index of the selected file in the uploaded_file_names list
    idx = uploaded_file_names.index(selected_file_name)
    
    # Check if at least one file is selected for chatting
    if selected_file_name:
        # Chatbot functionality
        st.markdown("<h3 style='text-align: center;'>CHAT WITH BOT</h3>", unsafe_allow_html=True)

        # Define function to generate response from user input
        def generate_response(input_text, dataframes):
            response = ""
            if input_text:
                sdf = SmartDataframe(dataframes[idx], config={"llm": llm})
                with st.spinner('Generating response...'):
                    response += "**Response for Selected File:**\n"
                    response += sdf.chat(input_text) + "\n"
                    
                    # Save the current figure to a temporary file
                    temp_chart_path = "temp_chart.png"
                    plt.savefig(temp_chart_path)
                    
                    # Display the graph if it exists
                    if os.path.exists(temp_chart_path):
                        st.image(temp_chart_path)
                        os.remove(temp_chart_path)  # Remove the temporary chart file after displaying

            else:
                st.warning('Enter input_text')

            return response

        # Container for chat history and download button
        response_container = st.container()
        download_button_container = st.container()

        # Container for text box
        input_container = st.container()

        with input_container:
            # Create a form for user input
            with st.form(key='my_form', clear_on_submit=True):
                user_input = st.text_area("**You:**", key='input', height=st.session_state.get('input_height', 50))
                st.session_state['input_height'] = len(user_input.split('\n')) * 20  # Adjust height based on input length
                submit_button = st.form_submit_button(label='Send')

            if submit_button and user_input:
                # If user submits input, generate response and store input and response in session state variables
                try:
                    response = generate_response(user_input, dataframes)
                    st.session_state['past'].append(user_input)
                    st.session_state['generated'].append(response)
                except Exception as e:
                    st.error("An error occurred: {}".format(e))

        if st.session_state['generated']:
            # Display chat history in a container
            with response_container:
                st.markdown("---")  # Adding a border line
                for i in range(len(st.session_state['generated'])):
                    message(st.session_state["past"][i], is_user=True, key=str(i) + '_user')
                    message(st.session_state["generated"][i], key=str(i))

            # Add a download button for chat history
            with download_button_container:
                download_button_clicked = False  # Flag to track whether download button has been clicked
                download_button = st.button("**Download Chat History**")

                if download_button and not download_button_clicked:
                    download_button_clicked = True  # Set flag to indicate button has been clicked
                    # Save chat history to a Word document
                    chat_history_file_path = "chat_history.docx"
                    document = Document()
                    document.add_heading('Chat History', level=1)
                    document.add_paragraph()

                    for user_input, generated_response in zip(st.session_state["past"], st.session_state["generated"]):
                        user_paragraph = document.add_paragraph()
                        user_paragraph.add_run("You: ").bold = True
                        user_paragraph.add_run(user_input).font.size = Pt(12)

                        bot_paragraph = document.add_paragraph()
                        bot_paragraph.add_run("Bot: ").bold = True
                        bot_paragraph.add_run(generated_response).font.size = Pt(12)

                        # Add some space between user and bot responses
                        document.add_paragraph()

                    document.save(chat_history_file_path)

                    # Provide a download link for the Word document
                    st.markdown(f"[Download Chat History](sandbox:/path/{chat_history_file_path})", unsafe_allow_html=True)